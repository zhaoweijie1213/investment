import json
import re
import shutil
from collections import defaultdict
from datetime import datetime
from html import unescape
from pathlib import Path
from urllib.parse import urlparse
from urllib.request import Request, urlopen
import ssl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"F:\Git\investment")
RAW_INPUT = ROOT / "docs" / "raw" / "lqy-weibo" / "all-2026-raw.json"
OUT_ROOT = ROOT / "docs" / "卢麒元微博"
RAW_DIR = OUT_ROOT / "raw"
IMG_DIR = OUT_ROOT / "images"
WORD_DIR = OUT_ROOT / "word"

INVESTMENT_KEYWORDS = [
    "投资", "资本", "金融", "财政", "货币", "美元", "人民币", "港币", "日元", "美债", "国债",
    "利率", "加息", "降息", "通胀", "恶通", "实质负利率", "汇率", "股", "股票", "A股",
    "港股", "美股", "日股", "黄金", "白银", "金银", "铜", "石油", "原油", "油价", "油气",
    "天然气", "粮", "粮食", "能源", "化工", "AI", "Token", "稳定币", "电币", "绿电",
    "短股长金", "金转油", "油粮", "向心坍缩", "速冻", "资本三流", "美联储", "沃什",
    "财政部", "关税", "债务", "房地产", "资产", "泡沫", "国储", "战略资源", "霍尔木兹",
    "伊朗", "日本", "套息",
]

ASSET_KEYWORDS = {
    "黄金/白银": ["黄金", "白银", "金银", "实质负利率"],
    "原油/能源": ["石油", "原油", "油价", "油气", "能源", "天然气", "化工", "霍尔木兹", "伊朗", "油粮"],
    "美元/美债": ["美元", "美债", "美国国债", "财政", "债务", "美联储", "沃什", "利率"],
    "人民币/港币": ["人民币", "港币", "绿电", "电币"],
    "日元/日本": ["日元", "日本", "套息", "日股"],
    "股票/风险资产": ["股票", "A股", "港股", "美股", "股市", "AI", "泡沫"],
    "粮食/必需品": ["粮", "粮食", "食品"],
    "Token/稳定币": ["Token", "稳定币", "数字资产"],
}


def strip_html(text: str) -> str:
    text = text or ""
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return unescape(text).replace("\u200b", "").strip()


def parse_weibo_date(value: str) -> datetime:
    return datetime.strptime(value, "%a %b %d %H:%M:%S %z %Y")


def contains_investment(text: str) -> bool:
    return any(k.lower() in text.lower() for k in INVESTMENT_KEYWORDS)


def asset_tags(text: str) -> str:
    tags = []
    for name, keys in ASSET_KEYWORDS.items():
        if any(k.lower() in text.lower() for k in keys):
            tags.append(name)
    return "、".join(tags) if tags else "宏观/投资"


def safe_name(value: str) -> str:
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value)
    return value[:120]


def image_url(pic: dict) -> str:
    for key in ("large", "largest", "original", "bmiddle", "thumbnail"):
        item = pic.get(key)
        if isinstance(item, dict) and item.get("url"):
            return item["url"].replace("\\/", "/")
    return ""


def collect_pics(mblog: dict):
    pics = []
    for p in mblog.get("pics") or []:
        url = image_url(p)
        if url:
            pics.append(url)
    retweeted = mblog.get("retweeted_status") or {}
    for p in retweeted.get("pics") or []:
        url = image_url(p)
        if url:
            pics.append(url)
    return list(dict.fromkeys(pics))


def download_image(url: str, target: Path) -> bool:
    if target.exists() and target.stat().st_size > 0:
        return True
    try:
        headers = {"User-Agent": "Mozilla/5.0", "Referer": "https://m.weibo.cn/"}
        req = Request(url, headers=headers)
        with urlopen(req, timeout=20, context=ssl._create_unverified_context()) as response:
            content = response.read()
        if content:
            target.write_bytes(content)
            return True
    except Exception:
        return False
    return False


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [0.9, 1.1, 1.2, 2.4, 0.9]
    headers = ["日期", "类型", "涉及资产", "核心原文", "来源"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, True)
        set_cell_shading(cell, "E8EEF5")
        cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        values = [row["date"], row["kind"], row["assets"], row["short_text"], row["source"]]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)
            cells[i].width = Inches(widths[i])
    return table


def add_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_meta_paragraph(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + "：")
    r.bold = True
    p.add_run(value)


def build_doc(month: str, records):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(f"卢麒元微博投资相关整理（{month}）")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    add_meta_paragraph(doc, "来源", "微博移动端公开页面；因微博第20页触发验证码，本批次仅覆盖已抓取公开数据")
    add_meta_paragraph(doc, "整理范围", f"{month} 月内投资、宏观、资产、货币、能源相关微博")
    add_meta_paragraph(doc, "风险提示", "原文观点不等于可验证事实，不构成买卖建议；涉及价格、政策、战争、利率与库存的数据须使用最新公开数据复核。")

    doc.add_heading("1. 月度概览", level=1)
    by_asset = defaultdict(int)
    for rec in records:
        for tag in rec["assets"].split("、"):
            by_asset[tag] += 1
    summary = "；".join([f"{k} {v}条" for k, v in sorted(by_asset.items(), key=lambda x: -x[1])[:8]])
    doc.add_paragraph(f"本月已纳入 {len(records)} 条投资相关微博。主题分布：{summary or '暂无'}。")

    doc.add_heading("2. 微博索引", level=1)
    rows = []
    for rec in records:
        rows.append({
            "date": rec["date"],
            "kind": rec["kind"],
            "assets": rec["assets"],
            "short_text": rec["plain_text"][:120] + ("..." if len(rec["plain_text"]) > 120 else ""),
            "source": rec["url"],
        })
    add_table(doc, rows)

    doc.add_heading("3. 原文整理", level=1)
    for idx, rec in enumerate(records, 1):
        doc.add_heading(f"{idx}. {rec['date']} - {rec['kind']}", level=2)
        add_meta_paragraph(doc, "涉及资产", rec["assets"])
        add_meta_paragraph(doc, "来源链接", rec["url"])
        if rec["truncated"]:
            add_meta_paragraph(doc, "完整性提示", "该条在公开接口中疑似为截断长文，需登录后复核全文。")
        p = doc.add_paragraph()
        r = p.add_run("原文观点：")
        r.bold = True
        doc.add_paragraph(rec["plain_text"])
        if rec.get("repost_text"):
            p = doc.add_paragraph()
            r = p.add_run("转发原文：")
            r.bold = True
            doc.add_paragraph(rec["repost_text"])
        if rec["image_files"]:
            p = doc.add_paragraph()
            r = p.add_run("图片：")
            r.bold = True
            for image_file in rec["image_files"]:
                try:
                    doc.add_picture(str(image_file), width=Inches(4.8))
                    last = doc.paragraphs[-1]
                    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph(f"图片文件无法嵌入：{image_file}")

    doc.add_heading("4. 出处说明", level=1)
    doc.add_paragraph("每条微博来源链接均保留在正文。图片保存于同级 images 目录；如微博登录态不足导致图片未下载，正文保留原图 URL。")
    out = WORD_DIR / f"2026-{month}_卢麒元微博投资相关整理.docx"
    doc.save(out)
    return out


def main():
    OUT_ROOT.mkdir(parents=True, exist_ok=True)
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(RAW_INPUT, RAW_DIR / "all-2026-raw-partial.json")

    raw = json.loads(RAW_INPUT.read_text(encoding="utf-8"))
    records_by_month = defaultdict(list)
    all_records = []
    for item in raw:
        m = item["mblog"]
        dt = parse_weibo_date(m["created_at"])
        plain = strip_html(m.get("text", ""))
        repost = strip_html((m.get("retweeted_status") or {}).get("text", ""))
        joined = plain + "\n" + repost
        if not contains_investment(joined):
            continue
        month = dt.strftime("%m")
        mid = m.get("mid") or m.get("id")
        rec_dir = IMG_DIR / f"2026-{month}" / str(mid)
        rec_dir.mkdir(parents=True, exist_ok=True)
        image_files = []
        image_urls = collect_pics(m)
        for i, url in enumerate(image_urls, 1):
            ext = Path(urlparse(url).path).suffix.lower()
            if ext not in [".jpg", ".jpeg", ".png", ".webp", ".gif"]:
                ext = ".jpg"
            target = rec_dir / f"{i:02d}{ext}"
            if download_image(url, target):
                image_files.append(target)
        kind = "原创"
        if m.get("retweeted_status"):
            kind = "转发/评论"
        if m.get("isLongText") or "全文" in plain:
            truncated = True
        else:
            truncated = False
        rec = {
            "date": dt.strftime("%Y-%m-%d %H:%M"),
            "month": month,
            "id": mid,
            "kind": kind,
            "assets": asset_tags(joined),
            "plain_text": plain,
            "repost_text": repost,
            "url": f"https://m.weibo.cn/status/{mid}",
            "image_urls": image_urls,
            "image_files": image_files,
            "truncated": truncated,
        }
        records_by_month[month].append(rec)
        all_records.append({**rec, "image_files": [str(p) for p in image_files]})

    index_path = RAW_DIR / "investment-records-partial.json"
    index_path.write_text(json.dumps(all_records, ensure_ascii=False, indent=2), encoding="utf-8")

    outputs = []
    for month, records in sorted(records_by_month.items()):
        records.sort(key=lambda r: r["date"], reverse=True)
        outputs.append(str(build_doc(month, records)))

    manifest = {
        "status": "partial",
        "coverage": "2026-04-24 to 2026-06-22 from public mobile Weibo pages; page 20 required captcha/login.",
        "raw_records": len(raw),
        "investment_records": len(all_records),
        "documents": outputs,
    }
    (OUT_ROOT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()




